from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--spec", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    spec = json.loads(Path(args.spec).read_text(encoding="utf-8"))
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    title = spec.get("title")
    if title:
        document.add_heading(title, level=0)

    subtitle = spec.get("subtitle")
    if subtitle:
        document.add_paragraph(subtitle)

    for paragraph in spec.get("paragraphs", []):
        document.add_paragraph(paragraph)

    for section in spec.get("sections", []):
        heading = section.get("heading")
        if heading:
            document.add_heading(heading, level=1)
        for paragraph in section.get("paragraphs", []):
            document.add_paragraph(paragraph)
        for bullet in section.get("bullets", []):
            document.add_paragraph(str(bullet), style="List Bullet")

    for table_spec in spec.get("tables", []):
        title_text = table_spec.get("title")
        if title_text:
            document.add_heading(title_text, level=2)
        headers = table_spec.get("headers", [])
        rows = table_spec.get("rows", [])
        if not headers:
            continue
        table = document.add_table(rows=1, cols=len(headers))
        header_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            header_cells[index].text = str(header)
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                if index < len(cells):
                    cells[index].text = str(value)

    document.save(str(output_path))


if __name__ == "__main__":
    main()
